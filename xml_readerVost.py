from xml.dom import minidom
from docx import Document
import re
import os
from docx.shared import Inches
import base64
from tkinter import Tk, Button, Frame, Entry, END
from tkinter import *
from tkinter import filedialog

root = Tk()
root.resizable(width=False, height=False)
root.geometry('{}x{}'.format(350, 500))
root.title("XMLReader")



tab_useless = ["<![CDATA[<p>", "</p>]]>", "</p>", "<p>", "<em>", "</em>", "<span style=\"font-family: courier new,courier,monospace;\">", "</span>"
    , "<span style=\"background-color: #ffff00;\">",
               "<span style=\"font-family: courier new,courier,monospace;\">",
               "<strong>", "</strong>", "<span style=\"font-family: courier new,courier,monospace;\">",
               "<pre>", "</pre>", "/>",
               "<span style=\"text-decoration: underline;\">"]

def delete_useless(text):

    t1 = text
    for i in tab_useless:
        t2 = t1.replace(i,'')
        t1 = t2
    t1 = re.sub(r'<img [\w.=/";:\d,\' &*()^%$#@!]*((alt="")|([ ]height="\w*"[ ]width="\w*"))', '', t1)
    t1 = re.sub(r'width="\w*"[ ]height="\w*"', '', t1)
    t1 = re.sub(r'height="\w*"[ ]width="\w*"', '', t1)
    t1 = re.sub(r'width="\w*"', '', t1)
    t1 = re.sub(r'height="\w*"', '', t1)
    return t1

def find_pictures(text):
    try:
        start = '<img src="data:image/png;base64,'
        end = '"'
        return str((text.split(start))[1].split(end)[0])

    except:
        return 0

def convert_pictures(text):
    with open(pff, "wb") as fh:
        fh.write(base64.decodebytes(text))

pom_string =''

def get_filename():
    pom = filedialog.askopenfilename(title = "Choose your file", filetypes = (("xml files", "*.xml"), ("all files", "*.*")))
    var.set("Ścieżka: " + pom)
    global pom_string
    pom_string = pom
    root.update_idletasks()

pom_pom_string = ''

def stworz(v):
    s = v.get()
    global pom_pom_string
    pom_pom_string = s

pff = ''

def fun():
    file_name = pom_string
    userhome = os.path.expanduser('~')
    desktop = userhome + '/Desktop/'
    output_path = desktop
    output_name = str(pom_pom_string)
    output_name = output_name + '.docx'
    global pff
    pff = output_path + "pom_image.png"
    filepath = output_path + output_name
    document = Document()
    DOMTree = minidom.parse(file_name)
    cNodes = DOMTree.childNodes
    tab1 = []
    tab2 = []
    j = 0
    for i in cNodes[0].getElementsByTagName("questiontext"):
        # print("Pytanie {}".format(j))
        tab1.append(i.getElementsByTagName("text")[0].childNodes[0].toxml())
        # print(i.getElementsByTagName("text")[0].childNodes[0].toxml())

    for i in cNodes[0].getElementsByTagName("answer"):
        # print(i.getElementsByTagName("text")[0].childNodes[0].toxml())
        tab2.append(i.getElementsByTagName("text")[0].childNodes[0].toxml())

    tab3 = ['a) ', 'b) ', 'c) ', 'd) ']
    p, q, r = 0, 0, 0
    for x in tab1:
        # print(tab1[p])
        s = find_pictures(str(tab1[p]))
        if (s):
            k = bytes(s, 'utf-8')
            # print(s)
            convert_pictures(k)
            newstr = tab1[p]
            newstr = newstr.replace(s, '')
            newstr = delete_useless(newstr)  # -------------------------------
            newstr = str(p + 1) + '. ' + newstr
            paragraph = document.add_paragraph(newstr)
            document.add_picture(pff, width=Inches(3))
        else:
            newstr = delete_useless(str(tab1[p]))
            newstr = str(p + 1) + '. ' + newstr
            paragraph = document.add_paragraph(newstr)
        p += 1
        for a in range(4):
            # print(tab2[q])
            s = find_pictures(str(tab2[q]))
            if (s):
                k = bytes(s, 'utf-8')
                # print(s)
                convert_pictures(k)
                newstr2 = str(tab2[q])
                newstr2 = newstr.replace(s, '')
                newstr2 = delete_useless(newstr2)
                newstr2 = tab3[r] + newstr2
                p2 = document.add_paragraph(newstr2)
                document.add_picture(pff, width=Inches(3))
            else:
                newstr2 = delete_useless(str(tab2[q]))
                newstr2 = tab3[r] + newstr2
                p2 = document.add_paragraph(newstr2)
            q += 1
            r += 1
        # print('')
        r = 0
        p3 = document.add_paragraph('')

    document.save(filepath)

    # document.save(output_name)

    if os.path.isfile(pff):
        os.remove(pff)

    lz = Label(root, text = "STWORZONO PLIK !")
    lz.place(relx=.5, rely=.8, anchor=CENTER)

def prr(v):
    stworz(v)
    fun()

var = StringVar()
var.set("Nie wybrano pliku!")


b = Button(root, text='Wybierz plik .xml',command=get_filename,padx=30,pady=15, compound=CENTER)
b.place(relx=.5, rely=.1, anchor=CENTER)
l = Label(root, textvariable = var)
l.place(relx=.5, rely=.2, anchor=CENTER)
l3 = Label(root, text = "------------------------------------------------------------------------------------------------")
l3.place(relx=.5, rely=.25, anchor=CENTER)
l2 = Label(root, text = "Podaj nazwę pliku .docx")
l2.place(relx=.5, rely=.3, anchor=CENTER)
v = StringVar()
e = Entry(root, textvariable=v)
e.place(relx=.5, rely=.4, anchor=CENTER)
l2 = Label(root, text = "-| Plik zostanie utworzony na pulpicie |-")
l2.place(relx=.5, rely=.5, anchor=CENTER)
b = Button(root, text='Stwórz',command= lambda: prr(v), padx=30,pady=15, compound=CENTER)
b.place(relx=.5, rely=.7, anchor=CENTER)

root.mainloop()

