from xml.dom import minidom
from docx import Document
import re
import os
from docx.shared import Inches
import base64

def main():

    file_name = str(input("Xml file path(e.g. C:\\\\temp\\\\mydata.xml): "))
    output_name = str(input("Output docx file name(e.g. quiz1): "))
    output_name = output_name + '.docx'
    output_path = str(input("Output docx file path(e.g. C:\\\\temp\\\\): "))
    pff = output_path + "pom_image.png"
    filepath = output_path + output_name
    document = Document()

    tab_useless = ["<![CDATA[<p>", "</p>]]>", "</p>", "<p>", "<em>", "</em>", "<span style=\"font-family: courier new,courier,monospace;\">", "</span>"
                , "<span style=\"background-color: #ffff00;\">",
               "<span style=\"font-family: courier new,courier,monospace;\">",
                 "<strong>", "</strong>", "<span style=\"font-family: courier new,courier,monospace;\">",
               "<pre>", "</pre>", "/>",
               "<span style=\"text-decoration: underline;\">"]

    def delete_useless(text):

        t1 = text
        for i in tab_useless:
            t2 = t1.replace(i,'')
            t1 = t2
        t1 = re.sub(r'<img [\w.=/";:\d,\' &*()^%$#@!]*((alt="")|([ ]height="\w*"[ ]width="\w*"))', '', t1)
        t1 = re.sub(r'width="\w*"[ ]height="\w*"', '', t1)
        t1 = re.sub(r'height="\w*"[ ]width="\w*"', '', t1)
        t1 = re.sub(r'width="\w*"', '', t1)
        t1 = re.sub(r'height="\w*"', '', t1)
        return t1

    def find_pictures(text):
        try:
            start = '<img src="data:image/png;base64,'
            end = '"'
            return str((text.split(start))[1].split(end)[0])

        except:
            return 0

    def convert_pictures(text):
        with open(pff, "wb") as fh:
            fh.write(base64.decodebytes(text))

    DOMTree = minidom.parse(file_name)
    cNodes = DOMTree.childNodes
    tab1 =[]
    tab2 =[]
    j=0
    for i in cNodes[0].getElementsByTagName("questiontext"):
        #print("Pytanie {}".format(j))
        tab1.append(i.getElementsByTagName("text")[0].childNodes[0].toxml())
        #print(i.getElementsByTagName("text")[0].childNodes[0].toxml())

    for i in cNodes[0].getElementsByTagName("answer"):
        #print(i.getElementsByTagName("text")[0].childNodes[0].toxml())
        tab2.append(i.getElementsByTagName("text")[0].childNodes[0].toxml())

    tab3 = ['a) ', 'b) ', 'c) ', 'd) ']
    p, q, r = 0, 0, 0
    for x in tab1:
        #print(tab1[p])
        s = find_pictures(str(tab1[p]))
        if(s):
            k = bytes(s, 'utf-8')
            #print(s)
            convert_pictures(k)
            newstr = tab1[p]
            newstr = newstr.replace(s, '')
            newstr = delete_useless(newstr)  #-------------------------------
            newstr = str(p + 1) + '. ' + newstr
            paragraph = document.add_paragraph(newstr)
            document.add_picture(pff, width=Inches(3))
        else:
            newstr = delete_useless(str(tab1[p]))
            newstr = str(p + 1) + '. ' + newstr
            paragraph = document.add_paragraph(newstr)
        p+=1
        for a in range(4):
            #print(tab2[q])
            s = find_pictures(str(tab2[q]))
            if (s):
                k = bytes(s, 'utf-8')
                #print(s)
                convert_pictures(k)
                newstr2 = str(tab2[q])
                newstr2 = newstr.replace(s, '')
                newstr2 = delete_useless(newstr2)
                newstr2 = tab3[r] + newstr2
                p2 = document.add_paragraph(newstr2)
                document.add_picture(pff, width=Inches(3))
            else:
                newstr2 = delete_useless(str(tab2[q]))
                newstr2 = tab3[r] + newstr2
                p2 = document.add_paragraph(newstr2)
            q+=1
            r+=1
        #print('')
        r=0
        p3 = document.add_paragraph('')


    document.save(filepath)

    #document.save(output_name)

    if os.path.isfile(pff):
        os.remove(pff)

    print("Docx file created!")
    i = input("Press ENTER to continue")

main()